import docx
import phonenumbers
from natasha import (
    Segmenter,
    NewsEmbedding,
    NewsNERTagger,
    Doc
)

emb = NewsEmbedding()
segmenter = Segmenter()
ner_tagger = NewsNERTagger(emb)


def anon_ner(text):
    result = ''
    doc = Doc(text)
    doc.segment(segmenter)
    doc.tag_ner(ner_tagger)
    result_temp = ''
    last = 0
    for span in doc.spans:
        if span.type == 'PER':
            result_temp += text[last:span.start]
            result_temp += 'ИМЯ'
        if span.type == 'ORG':
            result_temp += text[last:span.start]
            result_temp += 'ОРГАНИЗАЦИЯ'
        if span.type == 'LOC':
            result_temp += text[last:span.start]
            result_temp += 'ЛОКАЦИЯ'
        if span.type == 'PER' or span.type == 'ORG' or span.type == 'LOC':
            last = span.stop
    result_temp += text[last:]
    result = result_temp
    result_temp = ""
    last = 0
    countries = ['AZ', 'AM', 'BY', 'KZ', 'KG', 'MD', 'RU', 'TJ', 'TM', 'UZ', 'UA']
    for country in countries:
        for match in phonenumbers.PhoneNumberMatcher(result, country):
            result_temp += result[last:match.start]
            result_temp += 'ТЕЛЕФОН '
            last = match.end
    result_temp += result[last:]
    result = result_temp
    return result


def doc_style_in(doc):
    styles = []
    font_sizes = []
    for paragraph in doc.paragraphs:
        styles.append(paragraph.style)
        for run in paragraph.runs:
            font_sizes.append(run.font.size)
    styles_table = []
    font_sizes_table = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraphTab in cell.paragraphs:
                    styles_table.append(paragraphTab.style)
                    for run in paragraphTab.runs:
                        font_sizes_table.append(run.font.size)
    appearance = dict()
    appearance['styles'] = styles
    appearance['font_sizes'] = font_sizes
    appearance['styles_table'] = styles_table
    appearance['font_sizes_table'] = font_sizes_table
    return appearance


def doc_anon(in_path: str, out_name: str):
    doc = docx.Document(in_path)

    in_styles = doc_style_in(doc)
    styles = in_styles['styles']
    font_sizes = in_styles['font_sizes']
    styles_table = in_styles['styles_table']
    font_sizes_table = in_styles['font_sizes_table']

    for paragraph in doc.paragraphs:
        paragraph.text = anon_ner(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = anon_ner(cell.text)

    for numParagraphs in range(len(doc.paragraphs)):
        doc.paragraphs[numParagraphs].style = styles[numParagraphs]

    size_num = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            try:
                run.font.size = font_sizes[size_num]
            except IndexError:
                break
        size_num += 1

    size_num = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for numParagraphs in range(len(cell.paragraphs)):
                    cell.paragraphs[numParagraphs].style = styles_table[numParagraphs]
                for paragraphTab in cell.paragraphs:
                    for run in paragraphTab.runs:
                        try:
                            run.font.size = font_sizes_table[size_num]
                        except IndexError:
                            break
                    size_num += 1
    doc.save('./temp/' + out_name + '.docx')
